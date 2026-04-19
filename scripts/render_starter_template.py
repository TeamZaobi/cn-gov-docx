#!/usr/bin/env python3
"""
Render a starter DOCX for Chinese formal documents.

This script creates a scaffold document with practical typography defaults.
It is not a substitute for an organization's official template.
"""

from __future__ import annotations

import argparse
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


TYPE_CHOICES = (
    "briefing-report",
    "implementation-plan",
    "official-letter",
    "request",
    "faq",
    "meeting-minutes",
    "general-formal",
)

STRICTNESS_CHOICES = ("strict-official", "formal-business", "light-formal")


def set_run_font(run, font_name: str, size_pt: int, bold: bool = False) -> None:
    run.bold = bold
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    r_pr = run._element.get_or_add_rPr()
    east_asia = r_pr.rFonts
    if east_asia is None:
        east_asia = OxmlElement("w:rFonts")
        r_pr.append(east_asia)
    east_asia.set(qn("w:eastAsia"), font_name)


def set_paragraph_style(paragraph, first_line_chars: int = 2, fixed_line_pt: int = 28) -> None:
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Pt(first_line_chars * 16)
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(fixed_line_pt)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)


def configure_page(section, strictness: str) -> None:
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)

    if strictness == "strict-official":
        left_right = 2.8
    elif strictness == "formal-business":
        left_right = 2.6
    else:
        left_right = 2.54

    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(left_right)
    section.right_margin = Cm(left_right)


def get_font_profile(strictness: str) -> dict[str, tuple[str, int, bool]]:
    if strictness == "strict-official":
        return {
            "title": ("方正小标宋简体", 22, False),
            "body": ("仿宋_GB2312", 16, False),
            "heading_1": ("黑体", 16, False),
            "heading_2": ("楷体_GB2312", 16, False),
        }
    if strictness == "light-formal":
        return {
            "title": ("黑体", 22, True),
            "body": ("宋体", 15, False),
            "heading_1": ("黑体", 15, True),
            "heading_2": ("黑体", 15, False),
        }
    return {
        "title": ("黑体", 22, True),
        "body": ("仿宋_GB2312", 16, False),
        "heading_1": ("黑体", 16, False),
        "heading_2": ("楷体_GB2312", 16, False),
    }


def add_title(document: Document, title: str, profile: dict[str, tuple[str, int, bool]]) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    set_run_font(run, *profile["title"])
    paragraph.paragraph_format.space_after = Pt(20)


def add_body_paragraph(document: Document, text: str, profile: dict[str, tuple[str, int, bool]]) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_style(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, *profile["body"])


def add_heading(document: Document, text: str, level: int, profile: dict[str, tuple[str, int, bool]]) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    key = "heading_1" if level == 1 else "heading_2"
    set_run_font(run, *profile[key])
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(6)


def add_footer(section, unit_name: str) -> None:
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if unit_name:
        paragraph.add_run(unit_name)


def add_briefing_report(document: Document, profile) -> None:
    add_body_paragraph(document, "一、基本情况", profile)
    add_body_paragraph(document, "请在此填写背景、范围和当前状态。", profile)
    add_body_paragraph(document, "二、主要进展", profile)
    add_body_paragraph(document, "请分点说明目前形成的阶段成果。", profile)
    add_body_paragraph(document, "三、存在问题与风险", profile)
    add_body_paragraph(document, "请说明关键问题、约束条件及影响范围。", profile)
    add_body_paragraph(document, "四、下一步安排", profile)
    add_body_paragraph(document, "请写明后续任务、责任分工和时间安排。", profile)


def add_implementation_plan(document: Document, profile) -> None:
    add_body_paragraph(document, "一、建设背景与依据", profile)
    add_body_paragraph(document, "请在此说明建设背景、政策依据和现状问题。", profile)
    add_body_paragraph(document, "二、建设目标与范围", profile)
    add_body_paragraph(document, "请界定目标、服务对象、覆盖范围和边界。", profile)
    add_body_paragraph(document, "三、总体方案", profile)
    add_body_paragraph(document, "请概述总体架构、关键机制和运行方式。", profile)
    add_body_paragraph(document, "四、分阶段实施安排", profile)
    add_body_paragraph(document, "请按阶段写明上线、小范围验证、迭代优化和推广路径。", profile)
    add_body_paragraph(document, "五、保障措施", profile)
    add_body_paragraph(document, "请写明组织、制度、技术、数据和运维保障。", profile)


def add_official_letter(document: Document, profile, addressee: str) -> None:
    add_body_paragraph(document, f"{addressee}：", profile)
    add_body_paragraph(document, "请在此说明来函背景、沟通事项或答复内容。", profile)
    add_body_paragraph(document, "特此函复。", profile)


def add_request(document: Document, profile, addressee: str) -> None:
    add_body_paragraph(document, f"{addressee}：", profile)
    add_body_paragraph(document, "一、请示事项", profile)
    add_body_paragraph(document, "请明确写出拟请示或申请的具体事项。", profile)
    add_body_paragraph(document, "二、有关依据与必要性", profile)
    add_body_paragraph(document, "请说明依据、背景和必要性。", profile)
    add_body_paragraph(document, "三、拟采取的安排", profile)
    add_body_paragraph(document, "请写明拟实施的具体安排。", profile)
    add_body_paragraph(document, "妥否，请批示。", profile)


def add_faq(document: Document, profile) -> None:
    add_body_paragraph(document, "适用范围：请在此说明本说明面向的用户范围和使用边界。", profile)
    add_heading(document, "问题 1：这里填写常见问题", 1, profile)
    add_body_paragraph(document, "答：这里填写简明准确的回答。", profile)
    add_heading(document, "问题 2：这里填写常见问题", 1, profile)
    add_body_paragraph(document, "答：这里填写简明准确的回答。", profile)


def add_meeting_minutes(document: Document, profile) -> None:
    add_body_paragraph(document, "时间：", profile)
    add_body_paragraph(document, "地点：", profile)
    add_body_paragraph(document, "参会人员：", profile)
    add_body_paragraph(document, "一、会议议题", profile)
    add_body_paragraph(document, "请在此概述会议目的和议题。", profile)
    add_body_paragraph(document, "二、讨论要点", profile)
    add_body_paragraph(document, "请记录主要讨论内容。", profile)
    add_body_paragraph(document, "三、形成决定与后续安排", profile)
    add_body_paragraph(document, "请写明决定事项、责任人和时间要求。", profile)


def add_general_formal(document: Document, profile) -> None:
    add_body_paragraph(document, "一、背景与目的", profile)
    add_body_paragraph(document, "请在此填写文档背景和目标。", profile)
    add_body_paragraph(document, "二、主要内容", profile)
    add_body_paragraph(document, "请在此按层级展开主体内容。", profile)
    add_body_paragraph(document, "三、结语", profile)
    add_body_paragraph(document, "请在此填写总结性结语。", profile)


def add_trailing_signature(document: Document, unit_name: str, use_date: str, profile) -> None:
    if not unit_name and not use_date:
        return
    document.add_paragraph()
    if unit_name:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run(unit_name)
        set_run_font(run, *profile["body"])
    if use_date:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run(use_date)
        set_run_font(run, *profile["body"])


def render_document(args) -> Document:
    document = Document()
    configure_page(document.sections[0], args.strictness)
    profile = get_font_profile(args.strictness)

    add_title(document, args.title, profile)

    content_builders = {
        "briefing-report": lambda: add_briefing_report(document, profile),
        "implementation-plan": lambda: add_implementation_plan(document, profile),
        "official-letter": lambda: add_official_letter(document, profile, args.addressee or "有关单位"),
        "request": lambda: add_request(document, profile, args.addressee or "有关主管部门"),
        "faq": lambda: add_faq(document, profile),
        "meeting-minutes": lambda: add_meeting_minutes(document, profile),
        "general-formal": lambda: add_general_formal(document, profile),
    }
    content_builders[args.type]()

    add_trailing_signature(document, args.unit, args.date_text, profile)
    add_footer(document.sections[0], args.unit)
    return document


def parse_args():
    parser = argparse.ArgumentParser(description="Create a starter DOCX for Chinese formal documents.")
    parser.add_argument("--type", choices=TYPE_CHOICES, default="general-formal")
    parser.add_argument("--strictness", choices=STRICTNESS_CHOICES, default="formal-business")
    parser.add_argument("--title", required=True)
    parser.add_argument("--addressee", default="")
    parser.add_argument("--unit", default="")
    parser.add_argument("--date-text", default=str(date.today()))
    parser.add_argument("--output", required=True)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    output_path = Path(args.output).expanduser().resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = render_document(args)
    document.save(str(output_path))
    print(f"written: {output_path}")


if __name__ == "__main__":
    main()
